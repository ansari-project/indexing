import os
from pathlib import Path
import typer
from rich.console import Console
from rich.table import Table
from dotenv import load_dotenv
from .claude_converter import FiqhClaudeConverter

# Load environment variables from .env file
load_dotenv()

app = typer.Typer()
console = Console()


@app.command()
def convert(
    input_dir: Path = typer.Argument(
        Path("sample_input_data/fiqh_cards"),
        help="Input directory containing DOCX files"
    ),
    output_dir: Path = typer.Argument(
        Path("sample_output_data/fiqh_cards"),
        help="Output directory for JSON files"
    ),
    api_key: str = typer.Option(None, "--api-key", help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
):
    """Convert DOCX fiqh cards to JSON format using Claude Opus 4.1"""

    # Get API key from environment if not provided
    if not api_key:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            console.print("[red]Error: No API key provided. Set ANTHROPIC_API_KEY or use --api-key[/red]")
            raise typer.Exit(1)

    console.print(f"[bold blue]Converting fiqh cards using Claude Opus 4.1...[/bold blue]")
    console.print(f"Input directory: {input_dir}")
    console.print(f"Output directory: {output_dir}")

    if not input_dir.exists():
        console.print(f"[red]Error: Input directory {input_dir} does not exist[/red]")
        raise typer.Exit(1)

    converter = FiqhClaudeConverter(api_key=api_key)
    converter.process_directory(input_dir, output_dir)

    console.print("[green]✓ Conversion complete![/green]")


@app.command()
def test(
    api_key: str = typer.Option(None, "--api-key", help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
):
    """Test conversion on sample.docx file using Claude"""
    input_file = Path("sample_input_data/fiqh_cards/sample.docx")
    output_dir = Path("sample_output_data/fiqh_cards")

    # Get API key from environment if not provided
    if not api_key:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            console.print("[red]Error: No API key provided. Set ANTHROPIC_API_KEY or use --api-key[/red]")
            raise typer.Exit(1)

    if not input_file.exists():
        console.print(f"[red]Error: {input_file} does not exist[/red]")
        raise typer.Exit(1)

    console.print(f"[bold blue]Testing Claude conversion on {input_file.name}...[/bold blue]")

    converter = FiqhClaudeConverter(api_key=api_key)
    issues = converter.convert_docx_to_json(input_file)

    # Save output
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "sample_claude.json"

    import json
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(issues, f, ensure_ascii=False, indent=2)

    # Display summary
    table = Table(title=f"Converted {len(issues)} Issues with Claude")
    table.add_column("Issue #", style="cyan")
    table.add_column("Question", style="green")
    table.add_column("Opinions", style="yellow")

    for issue in issues[:5]:  # Show first 5 issues
        question = issue.get('question', '')
        table.add_row(
            str(issue.get('issue_number', 0)),
            question[:50] + "..." if len(question) > 50 else question,
            str(len(issue.get('opinions', [])))
        )

    console.print(table)
    console.print(f"\n[green]✓ Output saved to {output_file}[/green]")


@app.command()
def preview():
    """Preview how a table will be converted to markdown (without calling Claude)"""
    from docx import Document

    input_file = Path("sample_input_data/fiqh_cards/sample.docx")

    if not input_file.exists():
        console.print(f"[red]Error: {input_file} does not exist[/red]")
        raise typer.Exit(1)

    console.print(f"[bold blue]Previewing markdown conversion for {input_file.name}...[/bold blue]")

    # Don't need API key for preview
    converter = FiqhClaudeConverter()
    doc = Document(str(input_file))

    # Just show the first table
    if doc.tables:
        markdown = converter.table_to_markdown(doc.tables[0])
        console.print("\n[bold cyan]Table 1 as Markdown:[/bold cyan]")
        console.print(markdown)
    else:
        console.print("[yellow]No tables found in document[/yellow]")


if __name__ == "__main__":
    app()