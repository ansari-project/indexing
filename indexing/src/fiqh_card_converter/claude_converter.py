import json
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
import anthropic


class FiqhClaudeConverter:
    """Converter that uses Claude Opus 4.1 to extract structured data from fiqh cards."""

    # JSON schema for the expected output
    FIQH_ISSUE_SCHEMA = {
        "type": "object",
        "properties": {
            "issue_number": {
                "type": "integer",
                "description": "The issue number extracted from 'مسألة (X)'"
            },
            "question": {
                "type": "string",
                "description": "The main fiqh question being addressed"
            },
            "context": {
                "type": "string",
                "description": "Background and context of the disagreement (تحرير محل الخلاف)"
            },
            "opinions": {
                "type": "array",
                "description": "Different scholarly opinions",
                "items": {
                    "type": "object",
                    "properties": {
                        "position": {
                            "type": "string",
                            "description": "The opinion/position on the issue"
                        },
                        "scholars": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "List of scholars who hold this opinion"
                        }
                    }
                }
            },
            "disagreement_reason": {
                "type": "string",
                "description": "The reason for disagreement among scholars (سبب الخلاف)"
            },
            "evidence": {
                "type": "object",
                "description": "Evidence for different positions (الأدلة)",
                "additionalProperties": {"type": "string"}
            },
            "preferred_opinion": {
                "type": "string",
                "description": "The preferred/strongest opinion (الراجح)"
            },
            "practical_impact": {
                "type": "string",
                "description": "Practical impact of the ruling (ثمرة الخلاف)"
            },
            "references": {
                "type": "string",
                "description": "References for this issue (مراجع المسألة)"
            }
        }
    }

    def __init__(self, api_key: str = None):
        """Initialize with Anthropic API key."""
        self.client = anthropic.Anthropic(api_key=api_key)

    def table_to_markdown(self, table) -> str:
        """Convert a DOCX table to clean markdown format."""
        markdown_lines = []

        for row_idx, row in enumerate(table.rows):
            # Get unique cells (merged cells appear multiple times)
            unique_cells = []
            prev_text = None
            for cell in row.cells:
                if cell.text != prev_text:
                    unique_cells.append(cell.text.strip())
                    prev_text = cell.text

            if not unique_cells:
                continue

            # Format the row based on its content
            if len(unique_cells) >= 2:
                label = unique_cells[0]

                # Handle multiple columns (especially for opinions and evidence)
                if len(unique_cells) > 2:
                    markdown_lines.append(f"## {label}")
                    for i in range(1, len(unique_cells)):
                        markdown_lines.append(f"### Option {i}")
                        markdown_lines.append(unique_cells[i])
                        markdown_lines.append("")
                else:
                    markdown_lines.append(f"## {label}")
                    markdown_lines.append(unique_cells[1])
                    markdown_lines.append("")

        return "\n".join(markdown_lines)

    def extract_with_claude(self, markdown_text: str) -> Dict[str, Any]:
        """Use Claude Opus 4.1 to extract structured data from markdown."""

        prompt = f"""You are an expert in Islamic jurisprudence (fiqh). I will provide you with text from a fiqh card that contains information about a specific Islamic legal issue.

Please extract the information and return it as valid JSON matching this exact schema:

{json.dumps(self.FIQH_ISSUE_SCHEMA, indent=2, ensure_ascii=False)}

Here is the fiqh card content:

{markdown_text}

Important instructions:
1. Extract the issue number from text like "مسألة (1)" - the number should be an integer
2. The question is typically found in the first row after the issue number
3. For opinions (الأقوال ونسبتها), carefully separate different positions and their associated scholars
4. Scholars are often separated by "/" within a position
5. Evidence (الأدلة) may be marked with ٭ symbols - create separate evidence entries for each
6. Return ONLY valid JSON, no preamble or explanation
7. If a field is not present in the text, use an empty string for strings, empty array for arrays, or empty object for objects

JSON Output:"""

        # Make the API call to Claude Opus 4.1
        response = self.client.messages.create(
            model="claude-opus-4-1-20250805",
            max_tokens=16384,
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                },
                {
                    "role": "assistant",
                    "content": "{"  # Prefill to ensure JSON output
                }
            ]
        )

        # Get the response and ensure it starts with {
        json_text = "{" + response.content[0].text

        # Parse and return the JSON
        try:
            return json.loads(json_text)
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON from Claude: {e}")
            print(f"Response was: {json_text}")
            return {}

    def process_table(self, table) -> Dict[str, Any]:
        """Process a single table into structured data using Claude."""
        # Convert table to markdown
        markdown = self.table_to_markdown(table)

        # Extract structured data using Claude
        return self.extract_with_claude(markdown)

    def convert_docx_to_json(self, docx_path: Path) -> List[Dict[str, Any]]:
        """Convert a DOCX file to JSON format using Claude."""
        doc = Document(str(docx_path))
        issues = []

        for i, table in enumerate(doc.tables):
            print(f"Processing table {i+1}/{len(doc.tables)}...")
            try:
                issue = self.process_table(table)
                if issue:
                    issues.append(issue)
            except Exception as e:
                print(f"Error processing table {i+1}: {e}")
                continue

        return issues

    def process_directory(self, input_dir: Path, output_dir: Path):
        """Process all DOCX files in a directory."""
        output_dir.mkdir(parents=True, exist_ok=True)

        for docx_file in input_dir.glob("*.docx"):
            print(f"Processing {docx_file.name}...")

            try:
                issues = self.convert_docx_to_json(docx_file)

                # Save to JSON file
                output_file = output_dir / f"{docx_file.stem}.json"
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(issues, f, ensure_ascii=False, indent=2)

                print(f"  ✓ Converted to {output_file.name} ({len(issues)} issues)")

            except Exception as e:
                print(f"  ✗ Error: {e}")